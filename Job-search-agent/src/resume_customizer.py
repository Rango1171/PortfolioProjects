"""
Generates a tailored resume DOCX for each job using an LLM.
Supports Gemini, Groq (Llama / Gemma), and Mistral with automatic fallback.
"""

import logging
import re
import time
from pathlib import Path
from datetime import date

logger = logging.getLogger(__name__)

SYSTEM_PROMPT = (
    "You are an expert resume writer specializing in the European job market and ATS optimization. "
    "Your task is to tailor a candidate's resume to a specific job description.\n\n"
    "RULES:\n"
    "- NEVER invent skills, experience, or education the candidate does not have\n"
    "- Reorder and reword existing bullet points to emphasize JD-relevant experience\n"
    "- Rewrite the professional summary (2-3 sentences max) to mirror the job's language\n"
    "- Inject relevant ATS friendly keywords from the JD naturally into existing bullet points\n"
    "- Keep the exact same sections as the original resume, in the same order\n"
    "- Achievement-focused bullets: start with a strong action verb, include metrics where available\n"
    "- Maximum 4-5 bullets per role — be selective and concise to stay within 2 pages\n"
    "- Do NOT mention age, photo, marital status, nationality, or US state abbreviations\n"
    "- Format all locations as 'City, Country' (e.g. 'Chicago, USA' not 'Chicago, IL')\n"
    "- If the base resume has language skills, always include them\n"
    "- Target: exactly 2 pages — trim aggressively to stay within this limit\n\n"
    "OUTPUT FORMAT — follow exactly, no deviations:\n"
    "- First non-empty line: candidate full name only (copy exactly from base resume)\n"
    "- Second line: contact info joined by | (e.g. Chicago, USA | +1 (312) 394-0266 | email | LinkedIn | GitHub)\n"
    "- Section headings: ALL CAPS with no colon (e.g. PROFESSIONAL SUMMARY)\n"
    "- Job/education entry headers: prefix with ## (e.g. ## Company Name | Role Title | City, Country | Jan 2022 – Dec 2023)\n"
    "- Bullet points: each line starts with -  (hyphen + space)\n"
    "- Blank line between sections\n"
    "- Do NOT use markdown bold (**), italic, or code fences\n"
    "- Output ONLY the resume text — no commentary, no explanations, no preamble"
)


def _load_base_resume(base_resume_dir: str) -> str:
    base_path = Path(base_resume_dir)
    for ext, loader in [(".docx", _read_docx), (".pdf", _read_pdf), (".txt", _read_txt)]:
        for f in base_path.glob(f"*{ext}"):
            text = loader(f)
            if text.strip():
                logger.info(f"Loaded base resume from {f.name}")
                return text
    raise FileNotFoundError(
        f"No resume file found in {base_resume_dir}. "
        "Place a .docx, .pdf, or .txt resume file there."
    )


def _read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_pdf(path: Path) -> str:
    import pdfplumber
    with pdfplumber.open(str(path)) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _add_bottom_border(paragraph):
    """Add a thin bottom border under a paragraph (used for section headings)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _save_as_docx(text: str, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)

    def styled_para(content, size=10, bold=False, centered=False,
                    color=None, space_before=0, space_after=2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if centered:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(content)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        return p

    lines = [l for l in text.strip().split("\n")]
    non_empty_idx = 0  # counts only non-empty lines seen so far

    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # First non-empty line → candidate name (Heading: 16pt bold centered)
        if non_empty_idx == 0:
            styled_para(stripped, size=16, bold=True, centered=True,
                        space_before=0, space_after=3)
            non_empty_idx += 1
            continue

        # Second non-empty line → contact info (10pt centered)
        if non_empty_idx == 1 and ("|" in stripped or "@" in stripped):
            styled_para(stripped, size=10, bold=False, centered=True,
                        space_before=0, space_after=8)
            non_empty_idx += 1
            continue

        non_empty_idx += 1

        # Section headings: ALL CAPS, length > 3, not starting with digit or bullet
        if (stripped == stripped.upper()
                and len(stripped) > 3
                and not stripped[0].isdigit()
                and not stripped.startswith("-")):
            p = styled_para(stripped, size=12, bold=True,
                            color=RGBColor(0x1F, 0x49, 0x7D),
                            space_before=8, space_after=2)
            _add_bottom_border(p)
            continue

        # Job/education entry headers: ## prefix (subheading: 11pt bold)
        if stripped.startswith("## "):
            styled_para(stripped[3:].strip(), size=11, bold=True,
                        space_before=4, space_after=1)
            continue

        # Bullet points: 10pt, tight spacing
        if stripped.startswith(("- ", "• ", "* ", "· ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(re.sub(r"^[-•*·]\s+", "", stripped))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            continue

        # Regular content: 10pt
        styled_para(stripped, size=10, space_before=0, space_after=2)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _sanitize_filename(text: str) -> str:
    return re.sub(r"[^\w\-]", "_", text)[:40]


class ResumeCustomizer:
    """
    Stateful customizer. Reuse one instance per run — the LLM router and base
    resume text are loaded once and reused across all jobs.
    """

    def __init__(self, router, base_resume_dir: str, output_dir: str):
        self.router = router
        self.base_resume_text = _load_base_resume(base_resume_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _call_with_retry(self, prompt: str, max_retries: int = 3) -> str:
        """Call the LLM router; retry on transient (non-quota) failures."""
        last_err: Exception | None = None
        for attempt in range(max_retries):
            try:
                return self.router.generate(prompt, SYSTEM_PROMPT)
            except RuntimeError:
                raise  # all providers exhausted — surface immediately
            except Exception as e:
                last_err = e
                if attempt < max_retries - 1:
                    wait = 10 * (attempt + 1)
                    logger.warning(f"Customizer LLM error: {e}. Retrying in {wait}s...")
                    time.sleep(wait)
        raise RuntimeError(f"Customizer failed after {max_retries} attempts") from last_err

    def customize(self, job: dict) -> Path:
        """
        Generate a tailored resume for `job` and save it as a DOCX.
        Returns the path to the saved file.
        """
        company = _sanitize_filename(job.get("company", "company"))
        title = _sanitize_filename(job.get("title", "role"))
        today = date.today().isoformat()
        filename = f"{company}-{title}-{today}.docx"
        output_path = self.output_dir / filename

        jd_text = job.get("description", "")[:6000] or "(No description available)"

        prompt = (
            f"BASE RESUME:\n\n{self.base_resume_text}\n\n"
            f"---\n\n"
            f"JOB DESCRIPTION ({job.get('company', '')} — {job.get('title', '')}):\n\n"
            f"{jd_text}\n\n"
            "Produce the tailored resume now."
        )

        logger.info(f"Calling LLM to tailor resume for {company} / {title}...")
        # Free tier: 5 RPM. Wait 13s before each call to stay safely under the limit.
        time.sleep(13)
        tailored_text = self._call_with_retry(prompt)

        _save_as_docx(tailored_text, output_path)
        logger.info(f"Resume saved: {output_path.name}")
        return output_path
